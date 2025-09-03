import os
import uuid
import requests
from bson.objectid import ObjectId
import mysql.connector
import logging
from pymongo import MongoClient
from bson.errors import InvalidId
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import re
import pythoncom

logger = logging.getLogger('APP')
mongo_client = MongoClient('mongodb://localhost:27017/')
mongo_db = mongo_client['question_generator']

def add_text_to_doc(doc, text):
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('Section'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)
        elif line[0].isdigit() and '.' in line[:5]:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.bold = False
            run.font.size = Pt(11)
        elif line.startswith(('a)', 'b)', 'c)', 'd)')):
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(11)

def parse_exam_pattern_marks(pattern_text):
    """
    Parse exam pattern text to extract total marks and section-wise distribution.
    
    Args:
        pattern_text (str): Text content of the exam pattern.
    
    Returns:
        tuple: (total_marks, sections)
            - total_marks (int): Total marks for the question paper.
            - sections (list): List of dictionaries with section details (e.g., {'section': 'A', 'marks_per_question': 1, 'num_questions': 10}).
    """
    total_marks = 80  # Default total marks
    sections = []
    
    # Extract total marks
    total_marks_match = re.search(r'Total Marks:\s*(\d+)', pattern_text, re.IGNORECASE)
    if total_marks_match:
        total_marks = int(total_marks_match.group(1))
    
    # Extract section details (e.g., "Section A: 10 questions of 1 mark each")
    section_pattern = re.compile(r'Section\s+([A-Z]):\s*(\d+)\s*questions?\s*(?:of\s*(\d+)\s*mark)?', re.IGNORECASE)
    matches = section_pattern.findall(pattern_text)
    
    for match in matches:
        section = match[0]
        num_questions = int(match[1])
        marks_per_question = int(match[2]) if match[2] else 1
        sections.append({
            'section': section,
            'marks_per_question': marks_per_question,
            'num_questions': num_questions
        })
    
    if not sections:
        # Fallback structure if no sections are found
        sections = [
            {'section': 'A', 'marks_per_question': 1, 'num_questions': 10},
            {'section': 'B', 'marks_per_question': 3, 'num_questions': 5},
            {'section': 'C', 'marks_per_question': 5, 'num_questions': 5}
        ]
        total_marks = sum(d['marks_per_question'] * d['num_questions'] for d in sections)
    
    return total_marks, sections

def generate_question_paper(pattern_id, class_level, subject, mysql_config, upload_folder, output_folder, difficulty=None, selected_chapters=None):
    """
    Generate a question paper based on the exam pattern, chapter content, and difficulty level.
    
    Args:
        pattern_id (str): ID of the exam pattern in MySQL.
        class_level (str): Class level (e.g., '10').
        subject (str): Primary subject name (e.g., 'Science').
        mysql_config (dict): MySQL connection configuration.
        upload_folder (str): Path to the upload folder.
        output_folder (str): Path to the output folder.
        difficulty (str): Difficulty level ('easy', 'medium', 'hard').
        selected_chapters (list): List of selected chapters (e.g., ['Chapter 1', 'Physics:Motion']).
    
    Returns:
        dict: {'message': str, 'questions': str, 'docx_filename': str, 'pdf_filename': str} on success, or {'error': str} on failure.
    """
    try:
        # Step 1: Connect to MySQL
        conn = mysql.connector.connect(**mysql_config)
        cursor = conn.cursor(dictionary=True)

        # Step 2: Retrieve exam pattern from MySQL
        cursor.execute(
            """
            SELECT name, class_level, subject, subjects, sub_subject, chapters, mongo_id
            FROM exam_patterns
            WHERE id = %s
            """,
            (pattern_id,)
        )
        pattern = cursor.fetchone()
        if not pattern:
            cursor.close()
            conn.close()
            logger.warning(f"Exam pattern not found: ID={pattern_id}")
            return {'error': f'Exam pattern with ID {pattern_id} does not exist'}, 404

        subjects = pattern['subjects'].split(',') if pattern['subjects'] else [pattern['subject']] if pattern['subject'] else []
        subjects = [s.strip() for s in subjects]
        all_chapters = pattern['chapters'].split(',') if pattern['chapters'] else []
        all_chapters = [c.strip() for c in all_chapters]
        sub_subject = pattern['sub_subject']
        exam_type = pattern['name']
        mongo_id = pattern['mongo_id']
        logger.info(f"Retrieved exam pattern: ID={pattern_id}, name={exam_type}, subjects={subjects}, chapters={all_chapters}, mongo_id={mongo_id}")

        # Step 3: Filter chapters
        chapters = selected_chapters if selected_chapters else all_chapters
        chapters = [c.strip() for c in chapters if c.strip()]
        if not chapters:
            cursor.close()
            conn.close()
            logger.warning("No chapters provided for question paper generation")
            return {'error': 'At least one chapter must be selected'}, 400

        # Step 4: Fetch exam pattern text from MongoDB
        pattern_text = ""
        if mongo_id:
            try:
                pattern_doc = mongo_db['exam_pattern_combined'].find_one({'_id': ObjectId(mongo_id)})
                if pattern_doc and 'text' in pattern_doc:
                    pattern_text = pattern_doc['text']
                    logger.info(f"Retrieved exam pattern text from MongoDB: mongo_id={mongo_id}")
                else:
                    logger.warning(f"No text found in MongoDB for exam pattern: mongo_id={mongo_id}")
            except InvalidId:
                logger.error(f"Invalid MongoDB ObjectId: {mongo_id}")
                return {'error': 'Invalid exam pattern ID'}, 400
            except Exception as e:
                logger.error(f"Error retrieving exam pattern from MongoDB: mongo_id={mongo_id}, error={str(e)}")
                return {'error': f'Failed to retrieve exam pattern: {str(e)}'}, 500

        # Step 5: Parse exam pattern for marks
        total_marks, sections = parse_exam_pattern_marks(pattern_text)
        logger.info(f"Parsed exam pattern: total_marks={total_marks}, sections={sections}")

        # Step 6: Retrieve chapter mongo_ids from content table
        chapter_mongo_ids = []
        chapters_by_subject = {}
        for chapter in chapters:
            if not chapter:
                continue
            chap_subject = subject
            chap_name = chapter
            chap_sub_subject = None
            if ':' in chapter:
                parts = chapter.split(':')
                if len(parts) == 2:
                    chap_subject, chap_name = parts
                elif len(parts) == 3:
                    chap_subject, chap_sub_subject, chap_name = parts
                chap_subject = chap_subject.strip()
                chap_name = chap_name.strip()
                if chap_sub_subject:
                    chap_sub_subject = chap_sub_subject.strip()

            if chap_subject not in chapters_by_subject:
                chapters_by_subject[chap_subject] = []
            chapters_by_subject[chap_subject].append((chap_name, chap_sub_subject))

            query = """
                SELECT mongo_id, chapter
                FROM content
                WHERE class = %s AND subject = %s AND chapter = %s
            """
            params = [class_level, chap_subject, chap_name]
            if chap_sub_subject:
                query += " AND sub_subject = %s"
                params.append(chap_sub_subject)
            elif sub_subject and chap_subject == subject:
                query += " AND sub_subject = %s"
                params.append(sub_subject)

            cursor.execute(query, params)
            content = cursor.fetchone()
            if content and content['mongo_id']:
                chapter_mongo_ids.append((chap_subject, chap_sub_subject, content['mongo_id'], chap_name))
                logger.info(f"Found chapter in content: class={class_level}, subject={chap_subject}, sub_subject={chap_sub_subject or 'None'}, chapter={chap_name}, mongo_id={content['mongo_id']}")
            else:
                logger.warning(f"Chapter not found in content: class={class_level}, subject={chap_subject}, sub_subject={chap_sub_subject or 'None'}, chapter={chap_name}")

        cursor.close()
        conn.close()

        if not chapter_mongo_ids:
            logger.warning("No valid chapter content found in content table")
            return {'error': 'No content found for the selected chapters. Please upload relevant content.'}, 400

        # Step 7: Retrieve chapter content from MongoDB
        chapter_contents = []
        for chap_subject in chapters_by_subject:
            collection_name = f"class{class_level}_{chap_subject.lower().replace(' ', '_')}"
            try:
                mongo_ids = [mid for subj, sub_subj, mid, chap in chapter_mongo_ids if subj == chap_subject]
                if mongo_ids:
                    docs = mongo_db[collection_name].find({'_id': {'$in': [ObjectId(mid) for mid in mongo_ids]}})
                    for doc in docs:
                        if 'text' in doc:
                            for subj, sub_subj, mid, chap_name in chapter_mongo_ids:
                                if mid == str(doc['_id']):
                                    chapter_label = f"{subj}:{chap_name}" if not sub_subj else f"{subj}:{sub_subj}:{chap_name}"
                                    chapter_contents.append(f"Subject: {subj}{f', Sub-Subject: {sub_subj}' if sub_subj else ''}\nChapter: {chap_name}\nContent:\n{doc['text']}")
                                    break
                    logger.info(f"Retrieved {len(mongo_ids)} chapter contents from MongoDB: collection={collection_name}")
            except InvalidId:
                logger.error(f"Invalid MongoDB ObjectId in chapter_mongo_ids: {mongo_ids}")
                return {'error': 'Invalid chapter content ID'}, 400
            except Exception as e:
                logger.error(f"Error retrieving chapter content from MongoDB: collection={collection_name}, error={str(e)}")
                return {'error': f'Failed to retrieve chapter content: {str(e)}'}, 500

        if not chapter_contents:
            logger.warning("No chapter content available in MongoDB")
            return {'error': 'No content available for the selected chapters in MongoDB. Please upload relevant content.'}, 400

        # Step 8: Distribute marks equally across chapters
        num_chapters = len(chapters)
        if num_chapters == 0:
            logger.warning("No valid chapters found for mark distribution")
            return {'error': 'No valid chapters found for question generation'}, 400

        marks_per_chapter = total_marks // num_chapters
        remainder_marks = total_marks % num_chapters
        chapter_marks = {chapter: marks_per_chapter for chapter in chapters}
        for i, chapter in enumerate(chapters[:remainder_marks]):
            chapter_marks[chapter] += 1

        # Build mark distribution instruction
        mark_distribution_instruction = f"""
        - **Mark Distribution**:
          - Total marks: {total_marks}
          - Distribute questions equally across the {num_chapters} chapters to ensure balanced coverage.
          - Allocate approximately {marks_per_chapter} marks per chapter, with any remainder distributed to ensure the total is {total_marks}.
          - Chapter-wise mark allocation:
        """
        for chapter, marks in chapter_marks.items():
            mark_distribution_instruction += f"    - {chapter}: {marks} marks\n"

        # Step 9: Prepare prompt for Ollama
        combined_content = "\n\n".join(chapter_contents)
        difficulty_instructions = ""
        if difficulty == "easy":
            difficulty_instructions = """
            - **Difficulty: Easy**
            - Generate questions that directly test recall and basic understanding of the provided content.
            - Focus on straightforward facts, definitions, and concepts as presented in the book.
            - Examples: Direct questions like "What is...?", "Define...", or "List the...".
            - Ensure questions are simple and require minimal critical thinking.
            """
        elif difficulty == "medium":
            difficulty_instructions = """
            - **Difficulty: Medium**
            - Generate questions that require partial application of concepts from the provided content.
            - Combine recall with moderate problem-solving or interpretation, such as applying a concept to a simple scenario.
            - Examples: Questions like "Explain how...", "Give an example of...", or "Compare...".
            - Balance direct content usage with some analytical thinking.
            """
        elif difficulty == "hard":
            difficulty_instructions = """
            - **Difficulty: Hard**
            - Generate questions that require deep application, critical thinking, and problem-solving based on the provided content.
            - Focus on analytical, evaluative, or synthesis-based questions, such as solving complex problems or connecting concepts across chapters.
            - Examples: Questions like "Analyze...", "Evaluate the impact of...", or "Design a solution for...".
            - Ensure questions challenge students to think beyond the text.
            """
        else:
            logger.warning(f"Invalid difficulty level: {difficulty}")
            return {'error': 'Invalid difficulty level specified'}, 400

        section_instructions = ""
        for section in sections:
            section_instructions += f"""
            - Section {section['section']}: {section['num_questions']} questions, {section['marks_per_question']} mark(s) each
            """

        prompt = f"""
**Task**: Generate a question paper based on the provided exam pattern, content, and instructions.

**Exam Pattern Details**:
- Exam Type: {exam_type}
- Class: {class_level}
- Subject(s): {', '.join(subjects)}
- Sub-Subject: {sub_subject or 'N/A'}
- Chapters: {', '.join(chapters)}
- Difficulty: {difficulty}
- Pattern Specification:
{pattern_text}

**Content**:
{combined_content}

**Instructions**:
1. **Content Coverage**:
   - Ensure questions are based solely on the provided content.
   - Cover all specified subjects ({', '.join(subjects)}) and chapters ({', '.join(chapters)}).
   - Each chapter must be represented in the question paper.
2. **Mark Distribution**:
   {mark_distribution_instruction}
   - Ensure questions are distributed such that each chapter contributes approximately its allocated marks.
3. **Question Types**:
   - Follow the exam pattern structure:
     {section_instructions}
   - Include a mix of question types (e.g., multiple-choice, short answer, long answer) as specified in the exam pattern.
   - Distribute question types across chapters to maintain balance.
4. **Difficulty Level**:
   {difficulty_instructions}
5. **Formatting**:
   - Structure the question paper with clear sections (e.g., Section A, Section B) as per the exam pattern.
   - Number each question sequentially (e.g., Q1, Q2).
   - Specify marks for each question (e.g., [2 marks], [5 marks]).
   - Use plain text with consistent formatting for readability (e.g., use newlines between questions, bold section headers).
   - Indicate the chapter each question pertains to (e.g., "[Algebra] Question text").
6. **Constraints**:
   - Do not generate questions outside the provided content or chapters.
   - Avoid ambiguous or overly complex phrasing; ensure questions are clear and age-appropriate for Class {class_level}.
   - Output the question paper in plain text, properly formatted for direct use.

**Example Output Format**:
```
Section A: Multiple Choice Questions
1. [Algebra] [Question text] [1 mark]
   a) [Option 1]
   b) [Option 2]
   c) [Option 3]
   d) [Option 4]

Section B: Short Answer Questions
2. [Geometry] [Question text] [3 marks]

Section C: Long Answer Questions
3. [Physics:Motion] [Question text] [5 marks]
```
"""

        # Step 10: Call Ollama API
        ollama_url = "http://localhost:11434/api/generate"
        payload = {
            "model": "llama3:latest",
            "prompt": prompt,
            "stream": False
        }

        try:
            response = requests.post(ollama_url, json=payload, timeout=300)
            response.raise_for_status()
            result = response.json()
            if 'response' not in result:
                logger.error("Ollama response missing 'response' field")
                return {'error': 'Failed to generate questions: Invalid response from model'}, 500

            questions = result['response']
            logger.info(f"Generated question paper for pattern_id={pattern_id}, difficulty={difficulty}, subjects={subjects}, chapters={chapters}")

            # Step 11: Create DOCX
            doc = Document()
            p = doc.add_paragraph('RAMCO VIDYA MANDIR SR.SEC. SCHOOL, THAMARAIKULAM, ARIYALUR')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p = doc.add_paragraph(exam_type.upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
            p = doc.add_paragraph(f'CLASS: {class_level}\t\t\t\t\t\t\t\t\t\tMARKS: {total_marks}')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(11)
            p = doc.add_paragraph(f'SUBJECT: {", ".join(subjects)}\t\t\t\t\t\t\t\t\t\t\t\t\tTIME: 3Hrs')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(11)
            add_text_to_doc(doc, questions)
            docx_filename = f"question_paper_{uuid.uuid4().hex}_{difficulty}.docx"
            docx_filepath = os.path.join(output_folder, docx_filename)
            doc.save(docx_filepath)
            logger.info(f"Saved DOCX to {docx_filepath}")

            # Step 12: Convert to PDF
            try:
                pythoncom.CoInitialize()  # Initialize COM for this thread
                pdf_filename = docx_filename.replace('.docx', '.pdf')
                pdf_filepath = os.path.join(output_folder, pdf_filename)
                convert(docx_filepath, pdf_filepath)
                logger.info(f"Converted to PDF: {pdf_filepath}")
            except Exception as e:
                logger.error(f"Failed to convert DOCX to PDF: {str(e)}")
                return {'error': f'Failed to convert DOCX to PDF: {str(e)}. DOCX generated: {docx_filename}'}, 500
            finally:
                pythoncom.CoUninitialize()  # Clean up COM

            return {
                'message': 'Question paper generated successfully',
                'questions': questions,
                'docx_filename': docx_filename,
                'pdf_filename': pdf_filename
            }

        except requests.exceptions.RequestException as e:
            logger.error(f"Error calling Ollama API: {str(e)}")
            return {'error': f'Failed to generate questions: {str(e)}'}, 500

    except mysql.connector.Error as e:
        logger.error(f"MySQL error: {str(e)}")
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()
        return {'error': f'Database error: {str(e)}'}, 500
    except Exception as e:
        logger.exception(f"Error generating question paper: pattern_id={pattern_id}")
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()
        return {'error': f'Internal server error: {str(e)}'}, 500