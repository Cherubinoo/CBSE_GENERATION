import os
import re
import time
import logging
import mysql.connector
import ollama
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid

# Logging
logger = logging.getLogger('GENERATOR')
logger.setLevel(logging.INFO)
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
    logger.addHandler(handler)

# MySQL Database Connection
mysql_config = {
    'host': os.getenv('MYSQL_HOST', 'localhost'),
    'user': os.getenv('MYSQL_USER', 'root'),
    'password': os.getenv('MYSQL_PASSWORD', ''),
    'database': os.getenv('MYSQL_DB', 'education_db'),
    'connect_timeout': 10
}

def get_db_connection():
    return mysql.connector.connect(**mysql_config)

def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        logger.info(f"Read file: {file_path}")
        return content
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        return f"Error reading file: {str(e)}"

def generate_questions(class_level, subject, text, difficulty, chapters=None):
    # Filter text by chapters if provided
    if chapters:
        chapters_list = chapters.split(',') if chapters else []
        filtered_text = ""
        for chapter in chapters_list:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT content FROM content WHERE class = %s AND subject = %s AND chapter = %s", 
                          (class_level, subject, chapter))
            chapter_content = cursor.fetchall()
            cursor.close()
            conn.close()
            filtered_text += " ".join([row[0] for row in chapter_content if row[0]]) + " "
        text = filtered_text or text  # Fallback to original text if no chapter content

    # CBSE Class 10 English specific prompt
    if class_level == '10' and subject.lower() == 'english':
        passage_prompt = {
            "easy": """
Generate two reading comprehension passages (150-200 words each) for CBSE Class X English (Code 184, 2022-23), based on themes of hope and trust from *First Flight* (e.g., 'A Letter to God') or *Footprints Without Feet* (e.g., 'The Thief's Story'). For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]), totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, recall-based] [1]
(ii) [MCQ with 4 options, recall-based] [1]
(iii) [Short-answer, ~40 words, direct comprehension] [2]
(iv) [Completion task, simple] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, recall-based] [1]
...
""",
            "medium": """
Generate two reading comprehension passages (150-200 words each) for CBSE Class X English (Code 184, 2022-23), based on themes of struggle and perseverance from *First Flight* (e.g., 'Nelson Mandela: Long Walk to Freedom') or *Footprints Without Feet*. For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]) requiring inference, totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, inferential] [1]
(ii) [MCQ with 4 options, inferential] [1]
(iii) [Short-answer, ~40 words, analytical] [2]
(iv) [Completion task, inferential] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, inferential] [1]
...
""",
            "hard": """
Generate two reading comprehension passages (150-200 words each) for CBSE Class X English (Code 184, 2022-23), based on themes of ethical dilemmas and critical thinking from *First Flight* or *Footprints Without Feet* (e.g., 'The Thief's Story'). For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]) requiring critical thinking, totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, critical thinking] [1]
(ii) [MCQ with 4 options, critical thinking] [1]
(iii) [Short-answer, ~40 words, synthesis] [2]
(iv) [Completion task, complex] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, critical thinking] [1]
...
"""
        }
        question_prompt = {
            "easy": f"""
Based on the following CBSE Class X English textbook content (Code 184, 2022-23, *First Flight* or *Footprints Without Feet*): {text[:2000]}, generate questions for:
2. **Section B: Writing and Grammar (20 marks)**:
    - Grammar (10 marks): 10 tasks (e.g., fill-in-the-blank with verb forms, error correction, dialogue reporting) [1M each], based on Class 10 syllabus.
    - Writing (10 marks): 2 tasks (1 letter [5M], 1 analytical paragraph [5M], 120-150 words each, based on *First Flight* themes like hope or trust).
3. **Section C: Literature (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each]) from *First Flight* or *Footprints Without Feet*.
    - 4 short-answer questions (40-50 words, 3M each) on textbook chapters (*First Flight*, *Footprints Without Feet*).
    - 2 long-answer questions (100-120 words, 6M each) on textbook themes/characters (e.g., Lencho, Anil).
Questions should be simple, directly from the text, and formatted like a CBSE Class X English paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
""",
            "medium": f"""
Based on the following CBSE Class X English textbook content (Code 184, 2022-23, *First Flight* or *Footprints Without Feet*): {text[:2000]}, generate questions for:
2. **Section B: Writing and Grammar (20 marks)**:
    - Grammar (10 marks): 10 tasks (e.g., reporting dialogues, error correction requiring analysis, sentence transformation) [1M each], aligned with Class 10 syllabus.
    - Writing (10 marks): 2 tasks (1 letter [5M], 1 analytical paragraph [5M], 120-150 words each, analyzing *First Flight* themes like perseverance).
3. **Section C: Literature (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each], inference-based) from *First Flight* or *Footprints Without Feet*.
    - 4 short-answer questions (40-50 words, 3M each) requiring analysis of textbook chapters.
    - 2 long-answer questions (100-120 words, 6M each) requiring evaluation of textbook themes/characters.
Questions should involve inference and analysis, formatted like a CBSE Class X English paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
""",
            "hard": f"""
Based on the following CBSE Class X English textbook content (Code 184, 2022-23, *First Flight* or *Footprints Without Feet*): {text[:2000]}, generate questions for:
2. **Section B: Writing and Grammar (20 marks)**:
    - Grammar (10 marks): 10 complex tasks (e.g., sentence transformation, complex dialogue reporting, advanced error correction) [1M each], aligned with Class 10 syllabus.
    - Writing (10 marks): 2 tasks (1 letter [5M], 1 analytical paragraph [5M], 120-150 words each, evaluating *Footprints Without Feet* themes like ethical dilemmas).
3. **Section C: Literature (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each], critical thinking) from *First Flight* or *Footprints Without Feet*.
    - 4 short-answer questions (40-50 words, 3M each) requiring critical evaluation of textbook chapters.
    - 2 long-answer questions (100-120 words, 6M each) requiring deep synthesis of textbook themes/characters.
Questions should require critical thinking and synthesis, formatted like a CBSE Class X English paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
"""
        }
    else:
        # Generic prompt for other classes and subjects
        passage_prompt = {
            "easy": f"""
Generate two reading comprehension passages (150-200 words each) for CBSE Class {class_level} {subject}, based on general themes relevant to the subject syllabus. For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]), totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, recall-based] [1]
(ii) [MCQ with 4 options, recall-based] [1]
(iii) [Short-answer, ~40 words, direct comprehension] [2]
(iv) [Completion task, simple] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, recall-based] [1]
...
""",
            "medium": f"""
Generate two reading comprehension passages (150-200 words each) for CBSE Class {class_level} {subject}, based on themes requiring inference from the subject syllabus. For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]) requiring inference, totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, inferential] [1]
(ii) [MCQ with 4 options, inferential] [1]
(iii) [Short-answer, ~40 words, analytical] [2]
(iv) [Completion task, inferential] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, inferential] [1]
...
""",
            "hard": f"""
Generate two reading comprehension passages (150-200 words each) for CBSE Class {class_level} {subject}, based on themes requiring critical thinking from the subject syllabus. For each passage, create 4 questions (2 MCQs [1M each], 1 short-answer [2M], 1 completion task [1M]) requiring critical thinking, totaling 8 questions (10 marks per passage). Format as:
Section - A: Reading Skills (20 Marks)
1. Read the following text: [10M]
[Passage 1 text]
(i) [MCQ with 4 options, critical thinking] [1]
(ii) [MCQ with 4 options, critical thinking] [1]
(iii) [Short-answer, ~40 words, synthesis] [2]
(iv) [Completion task, complex] [1]
2. Read the following text: [10M]
[Passage 2 text]
(i) [MCQ with 4 options, critical thinking] [1]
...
"""
        }
        question_prompt = {
            "easy": f"""
Based on the following CBSE Class {class_level} {subject} content: {text[:2000]}, generate questions for:
2. **Section B: Core Concepts (20 marks)**:
    - 10 short tasks (e.g., definitions, basic problems, fill-in-the-blanks) [1M each], based on the syllabus.
    - 2 descriptive tasks (e.g., short explanations, problem-solving) [5M each], based on syllabus themes.
3. **Section C: Advanced Application (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each]) from the syllabus.
    - 4 short-answer questions (40-50 words, 3M each) on syllabus topics.
    - 2 long-answer questions (100-120 words, 6M each) on syllabus themes/concepts.
Questions should be simple, directly from the text, and formatted like a CBSE Class {class_level} paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
""",
            "medium": f"""
Based on the following CBSE Class {class_level} {subject} content: {text[:2000]}, generate questions for:
2. **Section B: Core Concepts (20 marks)**:
    - 10 short tasks (e.g., problem-solving, transformations, analytical tasks) [1M each], requiring inference from the syllabus.
    - 2 descriptive tasks (e.g., analytical explanations, problem-solving) [5M each], analyzing syllabus themes.
3. **Section C: Advanced Application (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each], inference-based) from the syllabus.
    - 4 short-answer questions (40-50 words, 3M each) requiring analysis of syllabus topics.
    - 2 long-answer questions (100-120 words, 6M each) requiring evaluation of syllabus themes/concepts.
Questions should involve inference and analysis, formatted like a CBSE Class {class_level} paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
""",
            "hard": f"""
Based on the following CBSE Class {class_level} {subject} content: {text[:2000]}, generate questions for:
2. **Section B: Core Concepts (20 marks)**:
    - 10 complex tasks (e.g., advanced problem-solving, complex transformations, critical analysis) [1M each], aligned with the syllabus.
    - 2 descriptive tasks (e.g., evaluative explanations, complex problem-solving) [5M each], evaluating syllabus themes.
3. **Section C: Advanced Application (40 marks)**:
    - 2 extract-based questions (5M each): 4 sub-questions each (2 MCQs/short-answers [1M each], 2 short-answers [1.5M each], critical thinking) from the syllabus.
    - 4 short-answer questions (40-50 words, 3M each) requiring critical evaluation of syllabus topics.
    - 2 long-answer questions (100-120 words, 6M each) requiring deep synthesis of syllabus themes/concepts.
Questions should require critical thinking and synthesis, formatted like a CBSE Class {class_level} paper. Return in plain text with sections separated by double newlines and questions numbered as in CBSE format (e.g., 3(i), 6(a)). Include marks in brackets (e.g., [1], [2]). Start numbering from 3 for Section B.
"""
        }
    
    try:
        logger.info(f"Generating questions for Class {class_level} {subject}, difficulty: {difficulty}")
        passage_response = ollama.generate(model="llama3", prompt=passage_prompt[difficulty])
        section_a_text = passage_response['response']
        logger.debug(f"Section A Raw Output: {section_a_text[:500]}...")

        question_response = ollama.generate(model="llama3", prompt=question_prompt[difficulty])
        sections_b_c_text = question_response['response']
        logger.debug(f"Sections B and C Raw Output: {sections_b_c_text[:500]}...")
        
        full_questions = f"{section_a_text}\n\n{sections_b_c_text}"
        logger.info("Questions generated successfully")
        return full_questions
    except Exception as e:
        logger.error(f"Error generating questions with Ollama: {str(e)}")
        return f"Error generating questions with Ollama: {str(e)}. Ensure 'ollama serve' is running and 'llama3' model is pulled."

def generate_pdf_and_docx(class_level, subject, questions, difficulty, output_folder='static/outputs', filename_prefix=""):
    base_filename = f"question_paper_class{class_level}_{subject.lower().replace(' ', '_')}_{difficulty}_{filename_prefix}"
    pdf_file = os.path.join(output_folder, f"{base_filename}.pdf")
    docx_file = os.path.join(output_folder, f"{base_filename}.docx")
    
    try:
        logger.debug(f"Processing questions for PDF/DOCX: {questions[:500]}...")
        
        section_a_match = re.search(r"Section\s*[-:]?\s*A.*?(?=Section\s*[-:]?\s*B|Section\s*[-:]?\s*C|$)", questions, re.DOTALL | re.IGNORECASE)
        section_b_match = re.search(r"Section\s*[-:]?\s*B.*?(?=Section\s*[-:]?\s*C|$)", questions, re.DOTALL | re.IGNORECASE)
        section_c_match = re.search(r"Section\s*[-:]?\s*C.*", questions, re.DOTALL | re.IGNORECASE)

        section_a = section_a_match.group(0).strip().splitlines() if section_a_match else []
        section_b = section_b_match.group(0).strip().splitlines() if section_b_match else []
        section_c = section_c_match.group(0).strip().splitlines() if section_c_match else []

        logger.debug(f"Section A (first 5 lines): {section_a[:5]}")
        logger.debug(f"Section B (first 5 lines): {section_b[:5]}")
        logger.debug(f"Section C (first 5 lines): {section_c[:5]}")

        if not section_a:
            logger.warning("Section A empty - using placeholder")
            section_a = ["Section - A: Reading Skills (20 Marks)", "1. Read the following text: [10M]", "(Placeholder passage)", "(i) Placeholder [1]"]
        if not section_b:
            logger.warning("Section B empty - using placeholder")
            section_b = ["Section - B: Core Concepts (20 Marks)", "3. Placeholder question [1]"]
        if not section_c:
            logger.warning("Section C empty - using placeholder")
            section_c = ["Section - C: Advanced Application (40 Marks)", "6. Placeholder question [5]"]

        # PDF Generation
        doc = SimpleDocTemplate(pdf_file, pagesize=A4, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(name='Title', fontSize=16, leading=20, alignment=1, spaceAfter=12)
        section_style = ParagraphStyle(name='Section', fontSize=14, leading=16, spaceBefore=12, spaceAfter=8, fontName='Helvetica-Bold')
        question_style = ParagraphStyle(name='Question', fontSize=12, leading=14, spaceAfter=6, leftIndent=10)
        instruction_style = ParagraphStyle(name='Instruction', fontSize=12, leading=14, spaceBefore=12, spaceAfter=6)
        table_style = TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('FONT', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTSIZE', (0,0), (-1,-1), 10),
        ])

        content = []
        
        content.append(Paragraph(f"CBSE Class {class_level} {subject} Question Paper - {difficulty.title()} (80 Marks)", title_style))
        content.append(Paragraph("Time allowed: 3 hours", styles['Italic']))
        content.append(Paragraph("Maximum Marks: 80", styles['Italic']))
        content.append(Spacer(1, 0.2*inch))
        
        content.append(Paragraph("<b>General Instructions:</b>", instruction_style))
        instructions = [
            "(i) This question paper comprises 11 questions. All questions are compulsory.",
            "(ii) The question paper contains THREE sections -",
            "     Section - A: Reading Skills",
            "     Section - B: Core Concepts" if class_level != '10' or subject.lower() != 'english' else "     Section - B: Writing and Grammar",
            "     Section - C: Advanced Application" if class_level != '10' or subject.lower() != 'english' else "     Section - C: Literature",
            "(iii) Attempt questions based on specific instructions for each part."
        ]
        for inst in instructions:
            content.append(Paragraph(inst, instruction_style))
        content.append(Spacer(1, 0.2*inch))

        for line in section_a:
            if re.match(r"Section\s*[-:]?\s*A", line, re.IGNORECASE):
                content.append(Paragraph(line, section_style))
            elif re.match(r"^\d+\.\s*Read the following text", line):
                content.append(Paragraph(line, question_style))
            elif re.match(r"^\(\w+\)\s*.*", line):
                content.append(Paragraph(line, question_style))
            else:
                content.append(Paragraph(line, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))

        for line in section_b:
            if re.match(r"Section\s*[-:]?\s*B", line, re.IGNORECASE):
                content.append(Paragraph(line, section_style))
            elif "Grammar (" in line or "Writing (" in line:
                content.append(Paragraph(line, section_style))
            elif "Core Concepts (" in line:
                content.append(Paragraph(line, section_style))
            elif re.match(r"^\d+\.\s*Complete any ten", line):
                content.append(Paragraph(line, question_style))
            elif "Error | Correction" in line:
                table_data = [line.split('|')[:2]]
                content.append(Table(table_data, colWidths=[2*inch, 2*inch], style=table_style))
            elif re.match(r"^\d+\.\s*.*\[5\]", line):
                content.append(Paragraph(line, question_style))
            elif re.match(r"^\(\w+\)\s*.*", line):
                content.append(Paragraph(line, question_style))
            else:
                content.append(Paragraph(line, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))

        for line in section_c:
            if re.match(r"Section\s*[-:]?\s*C", line, re.IGNORECASE):
                content.append(Paragraph(line, section_style))
            elif re.match(r"^\d+\.\s*Read the given extracts", line) or \
                 re.match(r"^\d+\.\s*Answer any four", line) or \
                 re.match(r"^\d+\.\s*Answer any one", line):
                content.append(Paragraph(line, question_style))
            elif re.match(r"^\(\w+\)\s*.*", line):
                content.append(Paragraph(line, question_style))
            else:
                content.append(Paragraph(line, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))

        try:
            doc.build(content)
            logger.info(f"PDF generated at {pdf_file}")
        except Exception as e:
            logger.error(f"Error building PDF: {str(e)}")
            return f"PDF build failed: {str(e)}", None

        docx = Document()
        docx.add_heading(f"CBSE Class {class_level} {subject} Question Paper - {difficulty.title()} (80 Marks)", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx.add_paragraph("Time allowed: 3 hours", style='Normal').runs[0].italic = True
        docx.add_paragraph("Maximum Marks: 80", style='Normal').runs[0].italic = True
        docx.add_paragraph()

        docx.add_heading("General Instructions:", level=2)
        for inst in instructions:
            p = docx.add_paragraph(inst, style='ListBullet')
            p.paragraph_format.left_indent = Inches(0.5)
        docx.add_paragraph()

        for line in section_a:
            if re.match(r"Section\s*[-:]?\s*A", line, re.IGNORECASE):
                docx.add_heading(line, level=2)
            elif re.match(r"^\d+\.\s*Read the following text", line):
                docx.add_paragraph(line)
            elif re.match(r"^\(\w+\)\s*.*", line):
                docx.add_paragraph(line, style='ListBullet')
            else:
                docx.add_paragraph(line)
        docx.add_paragraph()

        for line in section_b:
            if re.match(r"Section\s*[-:]?\s*B", line, re.IGNORECASE):
                docx.add_heading(line, level=2)
            elif "Grammar (" in line or "Writing (" in line or "Core Concepts (" in line:
                docx.add_heading(line, level=3)
            elif re.match(r"^\d+\.\s*Complete any ten", line):
                docx.add_paragraph(line)
            elif "Error | Correction" in line:
                table = docx.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Error'
                hdr_cells[1].text = 'Correction'
                row_cells = table.add_row().cells
                row_cells[0].text = line.split('|')[0].strip()
                row_cells[1].text = line.split('|')[1].strip() if len(line.split('|')) > 1 else ''
            elif re.match(r"^\d+\.\s*.*\[5\]", line):
                docx.add_paragraph(line, style='ListNumber')
            elif re.match(r"^\(\w+\)\s*.*", line):
                docx.add_paragraph(line, style='ListBullet')
            else:
                docx.add_paragraph(line)
        docx.add_paragraph()

        for line in section_c:
            if re.match(r"Section\s*[-:]?\s*C", line, re.IGNORECASE):
                docx.add_heading(line, level=2)
            elif re.match(r"^\d+\.\s*Read the given extracts", line) or \
                 re.match(r"^\d+\.\s*Answer any four", line) or \
                 re.match(r"^\d+\.\s*Answer any one", line):
                docx.add_paragraph(line, style='ListNumber')
            elif re.match(r"^\(\w+\)\s*.*", line):
                docx.add_paragraph(line, style='ListBullet')
            else:
                docx.add_paragraph(line)
        docx.add_paragraph()

        docx.save(docx_file)
        logger.info(f"Word document generated at {docx_file}")

        if os.path.exists(pdf_file) and os.path.exists(docx_file):
            return pdf_file, docx_file
        else:
            return "Error: PDF or Word document not generated.", None
    except Exception as e:
        logger.error(f"Error generating PDF or Word: {str(e)}")
        return f"Error generating PDF or Word: {str(e)}", None

def generate_question_paper(class_level, subject, exam_pattern_id, difficulty, file=None, upload_folder='static/uploads', output_folder='static/outputs'):
    try:
        logger.info(f"Generating question paper: class_level={class_level}, subject={subject}, exam_pattern_id={exam_pattern_id}, difficulty={difficulty}")
        
        if not exam_pattern_id or not class_level or not subject:
            raise ValueError("Class, subject, and exam pattern are required")
        
        if difficulty not in ['easy', 'medium', 'hard']:
            raise ValueError(f"Invalid difficulty level: {difficulty}")
        
        # Fetch exam pattern details
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT chapters, description FROM exam_patterns WHERE id = %s AND class_level = %s", 
                      (exam_pattern_id, class_level))
        pattern = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not pattern:
            raise ValueError(f"Exam pattern not found: ID={exam_pattern_id} for Class {class_level}")
        
        # Read uploaded file or use content from database
        text_content = ""
        if file and file.filename:
            ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
            if ext not in ['txt']:
                raise ValueError("Unsupported file type. Allowed: txt")
            if file.content_length > 10 * 1024 * 1024:
                raise ValueError("File size exceeds 10MB limit")
            filename = f"{uuid.uuid4().hex}_{file.filename}"
            filepath = os.path.join(upload_folder, filename)
            file.save(filepath)
            logger.info(f"File saved: {filepath}")
            text_content = read_text_file(filepath)
            if "Error" in text_content:
                raise ValueError(text_content)
        
        # Generate questions
        questions = generate_questions(class_level, subject, text_content, difficulty, chapters=pattern['chapters'])
        logger.debug(f"Generated Questions: {questions[:500]}...")
        if "Error" in questions:
            raise ValueError(questions)
        
        # Generate PDF and DOCX
        timestamp = str(int(time.time()))
        pdf_path, docx_path = generate_pdf_and_docx(class_level, subject, questions, difficulty, output_folder=output_folder, filename_prefix=timestamp)
        
        if isinstance(pdf_path, str) and "Error" in pdf_path:
            raise ValueError(pdf_path)
        
        if os.path.exists(pdf_path) and os.path.exists(docx_path):
            pdf_url = os.path.join("static", "outputs", os.path.basename(pdf_path))
            docx_url = os.path.join("static", "outputs", os.path.basename(docx_path))
            logger.info(f"Question paper generated: pdf={pdf_url}, docx={docx_url}")
            return {
                'message': 'Question paper generated successfully',
                'questions': questions,
                'pdf_url': pdf_url,
                'docx_url': docx_url
            }
        else:
            raise ValueError("File generation failed")
    except Exception as e:
        logger.error(f"Error in generate_question_paper: {str(e)}")
        raise ValueError(f"Error generating question paper: {str(e)}")