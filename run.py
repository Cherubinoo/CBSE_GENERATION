import os
import uuid
import logging
import time
import subprocess
import atexit
import signal
import threading
import sys
import requests
import json
import re
import mysql.connector
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from pymongo import MongoClient
from bson.objectid import ObjectId
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from pdf2image import convert_from_path
from pdf2image.exceptions import PDFPageCountError
import ollama

# Load environment variables from .env file
load_dotenv()

# --- Flask App Initialization ---
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'your_secret_key_here')

# --- Configuration ---
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['OUTPUT_FOLDER'] = 'static/outputs'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'png', 'jpg', 'jpeg', 'doc', 'docx', 'txt'}
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Embedding service configuration
EMBEDDING_SERVICE_URL = os.getenv('EMBEDDING_SERVICE_URL', 'http://localhost:5001/generate_embedding')
ollama_model = os.getenv('OLLAMA_MODEL', 'llama3')

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO,
                    format='[%(asctime)s] [%(levelname)s] [%(name)s] %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')
logger = logging.getLogger('FLASK_APP')

# --- Database Connections ---
# MongoDB for content, exam patterns, and staff
try:
    client = MongoClient(os.getenv('MONGODB_URI', 'mongodb://localhost:27017/'))
    mongo_db = client['document_management']
    staff_collection = mongo_db['staff']
    content_collection = mongo_db['content']
    exam_patterns_collection = mongo_db['exam_patterns']
    logger.info("MongoDB connected successfully.")
except Exception as e:
    logger.error(f"Error connecting to MongoDB: {str(e)}")
    sys.exit(1)

# MySQL for subjects and other frontend data
mysql_config = {
    'host': os.getenv('MYSQL_HOST', 'localhost'),
    'user': os.getenv('MYSQL_USER', 'root'),
    'password': os.getenv('MYSQL_PASSWORD', ''),
    'database': os.getenv('MYSQL_DB', 'education_db'),
    'connect_timeout': 10
}

def get_mysql_connection():
    try:
        connection = mysql.connector.connect(**mysql_config)
        logger.info("MySQL connected successfully.")
        return connection
    except mysql.connector.Error as e:
        logger.error(f"Error connecting to MySQL: {str(e)}")
        return None

# --- External Services Management (Embedding Service) ---
embedding_process = None
is_embedding_service_ready = threading.Event()

def log_subprocess_output(process):
    for line in iter(process.stdout.readline, b''):
        logger.info(f"Embedding service: {line.strip()}")
    for line in iter(process.stderr.readline, b''):
        logger.error(f"Embedding service error: {line.strip()}")

def start_embedding_service():
    global embedding_process
    logger.info("Attempting to start embedding service...")
    try:
        embedding_process = subprocess.Popen(
            ["python", "embedding.py"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1
        )
        stdout_thread = threading.Thread(target=log_subprocess_output, args=(embedding_process,))
        stdout_thread.daemon = True
        stdout_thread.start()

        timeout = 120
        start_time = time.time()
        while time.time() - start_time < timeout:
            try:
                response = requests.get(f"{os.getenv('EMBEDDING_SERVICE_URL', 'http://localhost:5001')}/health", timeout=5)
                if response.status_code == 200:
                    logger.info("Embedding service is up and running!")
                    is_embedding_service_ready.set()
                    return True
            except requests.exceptions.RequestException:
                pass
            time.sleep(1)
        logger.error("Embedding service failed to start within the timeout period.")
        return False
    except Exception as e:
        logger.error(f"Error starting embedding service subprocess: {str(e)}")
        if embedding_process:
            embedding_process.kill()
        return False

def terminate_embedding_service():
    global embedding_process
    if embedding_process and embedding_process.poll() is None:
        logger.info("Terminating embedding service...")
        embedding_process.terminate()
        embedding_process.wait(timeout=10)
        if embedding_process.poll() is None:
            embedding_process.kill()
        logger.info("Embedding service terminated.")

atexit.register(terminate_embedding_service)
signal.signal(signal.SIGTERM, lambda *args: terminate_embedding_service())
signal.signal(signal.SIGINT, lambda *args: terminate_embedding_service())

if not start_embedding_service():
    logger.critical("Failed to start embedding service. Exiting application.")
    sys.exit(1)

# --- Helper Functions from convert.py, process_pdf.py, generator.py, store.py ---

# From convert.py
poppler_path = r"C:\\poppler\\Library\\poppler-24.08.0\\Library\\bin"
ocr = PaddleOCR(lang='en', use_angle_cls=False, use_gpu=False, show_log=False)

def extract_text_from_pdf(filepath):
    try:
        images = convert_from_path(filepath, poppler_path=poppler_path)
        all_text = []
        for i, image in enumerate(images):
            result = ocr.ocr(image, cls=True)
            if result and result[0]:
                for line in result[0]:
                    text = line[-1][0]
                    all_text.append(text)
        return "\n".join(all_text)
    except Exception as e:
        logger.error(f"Error during PDF text extraction: {e}")
        return None

# From gpu_service.py/embedding.py
def get_embedding_from_gpu(text):
    try:
        response = requests.post(
            EMBEDDING_SERVICE_URL,
            json={"text": text},
            timeout=60
        )
        response.raise_for_status()
        response_data = response.json()
        if "error" in response_data:
            raise Exception(response_data["error"])
        return response_data.get('embedding')
    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching embedding from GPU service: {e}")
        return {"error": str(e)}

# From store.py
def store_to_mongo(metadata, text, embedding):
    try:
        required_keys = {'class', 'subject', 'resource_type', 'filename'}
        if not all(key in metadata for key in required_keys):
            logger.error(f"Missing required metadata keys: {required_keys - set(metadata.keys())}")
            return False

        collection_name = f"class{metadata['class']}_{metadata['subject']}".lower().replace(" ", "_")
        collection = mongo_db[collection_name]
        document = {
            "filename": metadata["filename"],
            "class": metadata["class"],
            "subject": metadata["subject"],
            "resource_type": metadata["resource_type"],
            "text": text,
            "embedding": embedding,
            "upload_date": datetime.utcnow()
        }
        result = collection.insert_one(document)
        logger.info(f"Stored in MongoDB → Collection: {collection_name}, _id: {result.inserted_id}")
        return True
    except Exception as e:
        logger.error(f"MongoDB insert failed: {e}")
        return False

# From process_pdf.py
def process_uploaded_pdf(filepath, form_data):
    try:
        text = extract_text_from_pdf(filepath)
        if not text:
            return {'status': 'error', 'message': 'Text extraction failed'}

        embeddings = get_embedding_from_gpu(text)
        if isinstance(embeddings, dict) and embeddings.get("error"):
            return {'status': 'error', 'message': f"Embedding failed: {embeddings['error']}"}

        metadata = {
            'class': form_data.get('class_level'),
            'subject': form_data.get('subject'),
            'resource_type': form_data.get('resource_type'),
            'chapter': form_data.get('chapter'),
            'original_filename': os.path.basename(filepath),
            'text': text,
            'embeddings': embeddings,
            'upload_date': datetime.utcnow()
        }
        result = content_collection.insert_one(metadata)
        content_id = str(result.inserted_id)

        return {
            'status': 'success',
            'message': 'PDF processed and stored successfully',
            'content_id': content_id
        }
    except Exception as e:
        logger.error(f"Error processing uploaded PDF: {e}")
        return {'status': 'error', 'message': f"Internal server error: {e}"}

# From generator.py
def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return None

def generate_questions(class_level, subject, text_content, difficulty, chapters=None):
    try:
        prompt = f"""
        You are an expert question paper generator for the CBSE curriculum.
        Generate 10 multiple-choice questions and 5 short-answer questions based on the following text content.
        Class: {class_level}
        Subject: {subject}
        Chapters: {', '.join(chapters) if chapters else 'All'}
        Difficulty: {difficulty}

        Content:
        {text_content}

        Output must be in a structured format with clear headings for each question type.
        """
        response = ollama.chat(model=ollama_model, messages=[{'role': 'user', 'content': prompt}])
        return response['message']['content']
    except Exception as e:
        logger.error(f"Error generating questions with Ollama: {e}")
        return f"Error generating questions: {e}"

def generate_pdf_and_docx(class_level, subject, questions, difficulty, output_folder, filename_prefix):
    # PDF Generation
    pdf_filename = f"{filename_prefix}_{class_level}_{subject}.pdf"
    pdf_path = os.path.join(output_folder, pdf_filename)
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Question', fontSize=12, leading=14))
    story = []
    story.append(Paragraph(f"Question Paper - Class {class_level} - {subject}", styles['h1']))
    story.append(Paragraph(f"Difficulty: {difficulty}", styles['h2']))
    story.append(Spacer(1, 0.2 * inch))

    lines = questions.split('\n')
    for line in lines:
        if line.strip():
            story.append(Paragraph(line, styles['Question']))
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)

    # DOCX Generation
    docx_filename = f"{filename_prefix}_{class_level}_{subject}.docx"
    docx_path = os.path.join(output_folder, docx_filename)
    document = Document()
    document.add_heading(f"Question Paper - Class {class_level} - {subject}", 0)
    document.add_paragraph(f"Difficulty: {difficulty}")

    for line in questions.split('\n'):
        if line.strip():
            document.add_paragraph(line)

    document.save(docx_path)

    return pdf_path, docx_path

# --- Routes ---
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def login_required(f):
    def wrapper(*args, **kwargs):
        if 'admin' not in session:
            # Check if it's an AJAX request
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'error': 'Unauthorized access. Please log in.'}), 401
            # If not an AJAX request, redirect
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    wrapper.__name__ = f.__name__
    return wrapper

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if username == 'admin' and password == 'admin123':
            session['admin'] = True
            logger.info("Admin user logged in successfully")
            return redirect(url_for('admin_dashboard'))
        else:
            logger.warning("Failed admin login attempt")
            return render_template('admin_login.html', error="Invalid Credentials")
    return render_template('admin_login.html')

@app.route('/admin_logout')
def admin_logout():
    session.pop('admin', None)
    logger.info("Admin user logged out")
    return redirect(url_for('admin_login'))

@app.route('/admin_dashboard')
@login_required
def admin_dashboard():
    return render_template('admin_dashboard.html')

@app.route('/view_staff')
@login_required
def view_staff():
    try:
        staff_list = list(staff_collection.find())
        return render_template('view_staff.html', staff_list=staff_list)
    except Exception as e:
        logger.error(f"Error viewing staff: {str(e)}")
        return "Error viewing staff", 500

@app.route('/edit_staff/<id>', methods=['GET', 'POST'])
@login_required
def edit_staff(id):
    try:
        if request.method == 'POST':
            name = request.form['name']
            role = request.form['role']
            email = request.form['email']
            phone = request.form['phone']
            staff_collection.update_one(
                {'_id': ObjectId(id)},
                {'$set': {'name': name, 'role': role, 'email': email, 'phone': phone}}
            )
            logger.info(f"Staff updated: _id={id}, name={name}")
            return redirect(url_for('view_staff'))
        staff = staff_collection.find_one({'_id': ObjectId(id)})
        if not staff:
            logger.warning(f"Staff not found: _id={id}")
            return redirect(url_for('view_staff'))
        return render_template('edit_staff.html', staff=staff)
    except Exception as e:
        logger.error(f"Error editing staff: {str(e)}")
        return redirect(url_for('view_staff'))

@app.route('/delete_staff/<id>', methods=['GET'])
@login_required
def delete_staff(id):
    try:
        staff = staff_collection.find_one({'_id': ObjectId(id)})
        if not staff:
            logger.warning(f"Staff not found: _id={id}")
            return redirect(url_for('view_staff'))
        staff_collection.delete_one({'_id': ObjectId(id)})
        logger.info(f"Staff deleted: _id={id}, name={staff['name']}")
        return redirect(url_for('view_staff'))
    except Exception as e:
        logger.error(f"Error deleting staff: {str(e)}")
        return redirect(url_for('view_staff'))

@app.route('/upload_content', methods=['GET', 'POST'])
@login_required
def upload_content():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            form_data = request.form.to_dict()
            result = process_uploaded_pdf(filepath, form_data)
            
            if result['status'] == 'success':
                return jsonify({'message': 'File uploaded and processed successfully.'})
            else:
                return jsonify({'error': result['message']}), 500
    return render_template('upload_content.html')

@app.route('/create_exam_pattern', methods=['GET', 'POST'])
@login_required
def create_exam_pattern():
    if request.method == 'POST':
        try:
            pattern = {
                'class_level': request.form['class_level'],
                'exam_type': request.form['exam_type'],
                'subjects': request.form.getlist('subjects'),
                'chapters': request.form.getlist('chapters'),
                'difficulty': request.form['difficulty'],
                'date_created': datetime.utcnow()
            }
            result = exam_patterns_collection.insert_one(pattern)
            logger.info(f"New exam pattern created: {result.inserted_id}")
            return jsonify({'message': 'Exam pattern created successfully'})
        except Exception as e:
            logger.error(f"Error creating exam pattern: {e}")
            return jsonify({'error': f'Failed to create exam pattern: {e}'}), 500
    return render_template('create_exam_pattern.html')

@app.route('/edit_exam_pattern', methods=['GET'])
@login_required
def edit_exam_pattern():
    return render_template('edit_exam_pattern.html')

@app.route('/get_exam_patterns/<class_level>')
@login_required
def get_exam_patterns(class_level):
    try:
        patterns = list(exam_patterns_collection.find({'class_level': class_level}))
        for pattern in patterns:
            pattern['_id'] = str(pattern['_id'])
        return jsonify(patterns)
    except Exception as e:
        logger.error(f"Error fetching exam patterns: {e}")
        return jsonify({'error': f'Failed to fetch exam patterns: {e}'}), 500

@app.route('/update_exam_pattern/<pattern_id>', methods=['POST'])
@login_required
def update_exam_pattern(pattern_id):
    try:
        updates = request.form.to_dict()
        if 'subjects' in updates:
            updates['subjects'] = updates['subjects'].split(',')
        if 'chapters' in updates:
            updates['chapters'] = updates['chapters'].split(',')
            
        exam_patterns_collection.update_one({'_id': ObjectId(pattern_id)}, {'$set': updates})
        return jsonify({'message': 'Exam pattern updated successfully'})
    except Exception as e:
        logger.error(f"Error updating exam pattern: {e}")
        return jsonify({'error': f'Failed to update exam pattern: {e}'}), 500

@app.route('/delete_exam_pattern/<pattern_id>', methods=['POST'])
@login_required
def delete_exam_pattern(pattern_id):
    try:
        exam_patterns_collection.delete_one({'_id': ObjectId(pattern_id)})
        return jsonify({'message': 'Exam pattern deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting exam pattern: {e}")
        return jsonify({'error': f'Failed to delete exam pattern: {e}'}), 500

@app.route('/generate_question_paper', methods=['GET', 'POST'])
@login_required
def generate_question_paper():
    if request.method == 'POST':
        try:
            class_level = request.form['class_level']
            subject = request.form['subject']
            difficulty = request.form['difficulty']

            # Find relevant content from MongoDB
            content = content_collection.find_one({'class': class_level, 'subject': subject})
            if not content:
                return jsonify({'error': 'No content found for the selected class and subject.'}), 404

            text_content = content['text']
            questions = generate_questions(class_level, subject, text_content, difficulty)
            
            if "Error" in questions:
                return jsonify({'error': questions}), 500

            timestamp = str(int(time.time()))
            output_folder = app.config['OUTPUT_FOLDER']
            pdf_path, docx_path = generate_pdf_and_docx(class_level, subject, questions, difficulty, output_folder, timestamp)

            pdf_url = url_for('static', filename=f"outputs/{os.path.basename(pdf_path)}")
            docx_url = url_for('static', filename=f"outputs/{os.path.basename(docx_path)}")

            return jsonify({
                'message': 'Question paper generated successfully',
                'questions': questions,
                'pdf_url': pdf_url,
                'docx_url': docx_url
            })
        except Exception as e:
            logger.error(f"Error generating question paper: {e}")
            return jsonify({'error': f'Failed to generate question paper: {e}'}), 500

    return render_template('generate_question_paper.html')

@app.route('/manage_subjects', methods=['GET'])
@login_required
def manage_subjects():
    return render_template('manage_subjects.html')

@app.route('/add_subject', methods=['POST'])
@login_required
def add_subject():
    conn = get_mysql_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.cursor()
        class_level = request.form['class_level']
        subject_name = request.form['subject_name']
        
        cursor.execute("INSERT INTO subjects (class_level, subject_name) VALUES (%s, %s)", (class_level, subject_name))
        conn.commit()
        return jsonify({'message': 'Subject added successfully'})
    except Exception as e:
        logger.error(f"Error adding subject: {e}")
        return jsonify({'error': 'Failed to add subject'}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/get_subjects/<class_level>')
@login_required
def get_subjects(class_level):
    conn = get_mysql_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM subjects WHERE class_level = %s", (class_level,))
        subjects = cursor.fetchall()
        return jsonify(subjects)
    except Exception as e:
        logger.error(f"Error fetching subjects: {e}")
        return jsonify({'error': 'Failed to fetch subjects'}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/update_subject/<subject_id>', methods=['POST'])
@login_required
def update_subject(subject_id):
    conn = get_mysql_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.cursor()
        subject_name = request.form['subject_name']
        class_level = request.form['class_level']
        cursor.execute("UPDATE subjects SET subject_name = %s, class_level = %s WHERE id = %s", (subject_name, class_level, subject_id))
        conn.commit()
        return jsonify({'message': 'Subject updated successfully'})
    except Exception as e:
        logger.error(f"Error updating subject: {e}")
        return jsonify({'error': 'Failed to update subject'}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/delete_subject/<subject_id>', methods=['POST'])
@login_required
def delete_subject(subject_id):
    conn = get_mysql_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM subjects WHERE id = %s", (subject_id,))
        conn.commit()
        return jsonify({'message': 'Subject deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting subject: {e}")
        return jsonify({'error': 'Failed to delete subject'}), 500
    finally:
        if conn and conn.is_connected():
            cursor.close()
            conn.close()

@app.route('/analytics')
@login_required
def analytics():
    return render_template('analytics.html')

if __name__ == '__main__':
    is_embedding_service_ready.wait()
    logger.info("Starting Flask application...")
    app.run(debug=True, host="0.0.0.0", port=5000, use_reloader=False)

